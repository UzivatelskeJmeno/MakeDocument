# Script to extract regions labeled "Úloha <number>" from a PDF and save them as images,
# then assemble those images into a Word (.docx) document.
# Requires: PyMuPDF (fitz), Pillow (PIL), python-docx, tkinter (standard lib)

#import fitz  # PyMuPDF: for reading PDF pages and rendering to images
from fitz import Rect as fitz_Rect
from fitz import open as fitz_open
#from PIL import Image  # Pillow (kept if further image processing is needed)
from docx import Document  # python-docx: for creating Word documents
from docx.shared import Inches  # helper to size images in the docx
import os
#import re # for regular expressions
from re import compile
#import tkinter as tk
from tkinter import filedialog # for file selection dialog

# Use tkinter file dialog to pick a PDF file (no visible root window)
filedialog.Tk().withdraw()  # Hide the initial file dialog root window
#root = tk.Tk()
#root.withdraw()
PDF_PATH = filedialog.askopenfilename(
    
    initialdir=os.getcwd(),
    title="Select PDF file",
    filetypes=[("PDF files", "*.pdf")]
)
# Exit if user canceled the dialog
if not PDF_PATH:
    print("No PDF file selected. Exiting.")
    exit()
print(f"Selected PDF: {PDF_PATH}")

# Output locations (adjust as needed)
OUTPUT_DIR = f"{PDF_PATH.removesuffix('.pdf')}//uloha_images"
DOCX_PATH = f"{PDF_PATH.removesuffix('.pdf')}.docx"
os.makedirs(OUTPUT_DIR, exist_ok=True)  # create output folder if missing
if not os.path.exists(DOCX_PATH):
    doc = Document()
    doc.save(DOCX_PATH)
    print(f"Created new DOCX file: {DOCX_PATH}")

# Regular expression to find blocks that start with "Úloha <number>"
ULOHA_PATTERN = compile(r"Úloha\s+\d+")
print(ULOHA_PATTERN)

MARGIN = 4  # pixels to extend the crop above the detected text block
OFFSET_X = 35  # extra pixels to extend crop left/right if needed
def extract_uloha_regions():
    """
    Scan the selected PDF and extract rectangular regions corresponding to
    "Úloha <n>" headings. Each region is rendered to a PNG image saved into OUTPUT_DIR.
    Returns a list of saved image file paths.
    """
    doc = fitz_open(PDF_PATH)  # open the PDF
    images = []

    # Iterate through all pages in the PDF
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        # Get text "blocks" which include bounding box coords and text
        blocks = page.get_text("blocks")
        uloha_blocks = []

        # Find all blocks that match the Úloha pattern
        for block in blocks:
            text = block[4]  # text content is at index 4 in the block tuple
            if ULOHA_PATTERN.match(text.strip()):
                uloha_blocks.append(block)
        
        # Sort matched blocks top-to-bottom using the y0 coordinate (block[1])
        uloha_blocks.sort(key=lambda b: b[1])
        #print(f"uloha_blocks: {uloha_blocks}\n")
        print(f"page {page_num+1}")
        
        # Crop each Úloha region: from current block y to next block y (or page bottom)
        for i in range(len(uloha_blocks)):
            x0, y0, _, _ = uloha_blocks[i][:4]  # x0,y0 are the top-left of the block
            y1 = (
                uloha_blocks[i + 1][1] if i + 1 < len(uloha_blocks)
                else page.rect.height  # if last block, go to page bottom
            )
            # Define rectangular area to render: full width, from slightly above y0 to y1
            rect = fitz_Rect(OFFSET_X, y0 - MARGIN, page.rect.width-OFFSET_X, y1)
            # Render the region at 300 DPI for good quality
            pix = page.get_pixmap(clip=rect, dpi=300)

            # Create a safe label from the block text, replace spaces, trim length
            label = uloha_blocks[i][4].strip().replace(" ", "_")
            label = label[:8]  # limit filename length to avoid excessively long names
            img_path = os.path.join(OUTPUT_DIR, f"{label}.png")
            print(img_path)
            
            # Save rendered image and add to list
            pix.save(img_path)
            images.append(img_path)

    return images

def create_docx(images):
    """
    Create a Word document and insert each image (one per Úloha region).
    Adds a few empty paragraphs between images for spacing.
    """
    doc = Document()
    for img_path in images:
        doc.add_paragraph()  # optional spacing before the picture
        doc.add_picture(img_path, width=Inches(6))  # insert image sized to 6 inches wide
        doc.add_paragraph()  # add extra spacing after each image
        bold = doc.add_paragraph(text="Author: ")  # line for author signature
        # Normal text
        bold.add_run("text")

        # Bold text
        #bold_run.bold = True
        bold.runs[0].bold = True
    doc.core_properties.comments = ""
    doc.core_properties.author = os.getlogin()




    doc.save(DOCX_PATH)

if __name__ == "__main__":
    print("Extracting Úloha regions...")
    images = extract_uloha_regions()
    print(f"Extracted {len(images)} problems.")
    print("Creating Word document...")
    create_docx(images)
    print(f"Saved as {DOCX_PATH}")
